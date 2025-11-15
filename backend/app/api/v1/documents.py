"""
Document Upload and Analysis API endpoints
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc
from typing import List, Optional
from uuid import UUID
import logging
import os
from pathlib import Path
import uuid as uuid_lib
from datetime import datetime
import mimetypes

from app.core.deps import get_db, get_current_active_user
from app.models.document_knowledge import DocumentKnowledge, DocumentType, DocumentSource
from app.models.project import Project
from app.models.user import User

router = APIRouter()
logger = logging.getLogger(__name__)

# Configure upload settings
UPLOAD_DIR = Path("uploads/documents")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


async def verify_project_access(
    project_id: UUID,
    current_user: User,
    db: AsyncSession,
) -> Project:
    """Verify that the current user has access to the project."""
    result = await db.execute(
        select(Project).where(Project.id == project_id)
    )
    project = result.scalar_one_or_none()

    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found"
        )

    # Verify user has access to the organisation
    from sqlalchemy import text
    access_check = await db.execute(
        text("""
            SELECT 1 FROM organisations o
            LEFT JOIN user_organisations uo ON o.id = uo.organisation_id
            WHERE o.id = :org_id
            AND (o.owner_id = :user_id OR uo.user_id = :user_id)
            LIMIT 1
        """),
        {"org_id": str(project.organisation_id), "user_id": str(current_user.id)}
    )

    if not access_check.fetchone():
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You don't have permission to access this project"
        )

    return project


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    project_id: UUID = Form(...),
    document_type: str = Form("requirement"),
    document_name: Optional[str] = Form(None),
    source: str = Form("upload"),
    description: Optional[str] = Form(None),
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a document (BRD, requirements, etc.) for analysis.
    Supports PDF, DOC, DOCX, TXT, MD formats.
    """
    # Verify project access
    await verify_project_access(project_id, current_user, db)

    # Validate file extension
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Read file content
    try:
        content = await file.read()
        file_size = len(content)

        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024}MB"
            )

        # Generate unique filename
        unique_id = str(uuid_lib.uuid4())
        safe_filename = f"{unique_id}_{file.filename}"
        file_path = UPLOAD_DIR / safe_filename

        # Save file
        with open(file_path, "wb") as f:
            f.write(content)

        # Convert string to enum values
        try:
            doc_type_enum = DocumentType[document_type.upper()] if document_type.upper() in DocumentType.__members__ else DocumentType.DOCUMENT
        except (KeyError, AttributeError):
            doc_type_enum = DocumentType.DOCUMENT

        try:
            source_enum = DocumentSource[source.upper()] if source.upper() in DocumentSource.__members__ else DocumentSource.FILE_UPLOAD
        except (KeyError, AttributeError):
            source_enum = DocumentSource.FILE_UPLOAD

        # Create document record
        document = DocumentKnowledge(
            project_id=project_id,
            created_by=current_user.id,
            document_name=document_name or file.filename,
            content="",  # Will be extracted by ingestion service
            document_type=doc_type_enum,
            source=source_enum,
            file_type=file_ext,
            content_length=file_size,
            meta_data={
                "original_filename": file.filename,
                "file_path": str(file_path),
                "file_size": str(file_size),
                "uploaded_by": current_user.email,
                "description": description,
                "upload_timestamp": datetime.utcnow().isoformat(),
            }
        )

        db.add(document)
        await db.commit()
        await db.refresh(document)

        # TODO: Trigger async document processing
        # This would:
        # 1. Extract text from the document
        # 2. Chunk the text
        # 3. Create embeddings
        # 4. Store in vector database
        # For now, we'll just store the document record

        logger.info(f"Document uploaded: {document.id} for project {project_id}")

        return {
            "success": True,
            "message": "Document uploaded successfully",
            "document": {
                "id": str(document.id),
                "title": document.document_name,
                "file_type": document.file_type,
                "file_size": document.meta_data.get("file_size", "0"),
                "uploaded_at": document.created_at.isoformat(),
                "status": "processing",
            }
        }

    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}"
        )


@router.get("/")
async def list_documents(
    project_id: UUID,
    document_type: Optional[str] = None,
    source: Optional[str] = None,
    is_active: bool = True,
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    List all documents for a project with optional filtering.
    """
    # Verify project access
    await verify_project_access(project_id, current_user, db)

    # Build query
    query = select(DocumentKnowledge).where(DocumentKnowledge.project_id == project_id)

    if document_type:
        query = query.where(DocumentKnowledge.document_type == document_type)
    if source:
        query = query.where(DocumentKnowledge.source == source)
    if is_active is not None:
        query = query.where(DocumentKnowledge.is_active == is_active)

    query = query.order_by(desc(DocumentKnowledge.created_at)).offset(skip).limit(limit)

    result = await db.execute(query)
    documents = result.scalars().all()

    return {
        "documents": [
            {
                "id": str(doc.id),
                "title": doc.document_name,
                "document_type": doc.document_type.value if doc.document_type else "",
                "source": doc.source.value if doc.source else "",
                "file_type": doc.file_type,
                "file_size": doc.meta_data.get("file_size", "0"),
                "uploaded_by": doc.meta_data.get("uploaded_by", ""),
                "created_at": doc.created_at.isoformat(),
                "chunk_count": doc.total_chunks,
                "is_active": doc.is_active,
            }
            for doc in documents
        ],
        "total": len(documents),
    }


@router.get("/{document_id}")
async def get_document(
    document_id: UUID,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """Get document details."""
    result = await db.execute(
        select(DocumentKnowledge).where(DocumentKnowledge.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify project access
    await verify_project_access(document.project_id, current_user, db)

    return {
        "id": str(document.id),
        "title": document.document_name,
        "description": document.meta_data.get("description", ""),
        "document_type": document.document_type.value if document.document_type else "",
        "source": document.source.value if document.source else "",
        "file_path": document.meta_data.get("file_path", ""),
        "file_type": document.file_type,
        "file_size": document.meta_data.get("file_size", "0"),
        "uploaded_by": document.meta_data.get("uploaded_by", ""),
        "created_at": document.created_at.isoformat(),
        "updated_at": document.updated_at.isoformat() if document.updated_at else None,
        "chunk_count": document.total_chunks,
        "embedding_model": document.qdrant_collection,
        "is_active": document.is_active,
        "meta_data": document.meta_data,
    }


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: UUID,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Delete a document.
    This will also delete the associated file and vector embeddings.
    """
    result = await db.execute(
        select(DocumentKnowledge).where(DocumentKnowledge.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify project access
    await verify_project_access(document.project_id, current_user, db)

    # Delete file if exists
    file_path = document.meta_data.get("file_path")
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            logger.warning(f"Failed to delete file: {e}")

    # TODO: Delete vector embeddings from Qdrant

    # Delete database record
    await db.delete(document)
    await db.commit()

    logger.info(f"Document deleted: {document_id}")
    return None


@router.post("/{document_id}/analyze")
async def analyze_document(
    document_id: UUID,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Analyze a document to extract requirements, features, and testable scenarios.
    Uses AI to comprehensively analyze the document content.
    """
    result = await db.execute(
        select(DocumentKnowledge).where(DocumentKnowledge.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify project access
    await verify_project_access(document.project_id, current_user, db)

    try:
        # Get document content
        file_path = document.meta_data.get("file_path")
        content = ""

        if file_path and os.path.exists(file_path):
            # Extract content from file
            try:
                import PyPDF2
                file_ext = Path(file_path).suffix.lower()

                if file_ext == '.pdf':
                    # Extract PDF content
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        content_parts = []
                        for page in reader.pages:
                            content_parts.append(page.extract_text())
                        content = '\n'.join(content_parts)
                elif file_ext in ['.txt', '.md']:
                    # Read text files
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_ext == '.docx':
                    # Extract DOCX content
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                else:
                    content = document.content or ""
            except Exception as e:
                logger.warning(f"Error extracting content from file: {e}")
                content = document.content or ""
        else:
            # Use stored content
            content = document.content or ""

        if not content or len(content) < 50:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document content is too short or empty to analyze"
            )

        # Analyze document using AI
        from app.services.document_analysis_service import get_document_analysis_service
        from app.services.ai_service import AIService

        ai_service = AIService()
        analysis_service = get_document_analysis_service(ai_service)

        analysis_result = await analysis_service.analyze_document_content(
            content=content,
            document_type=document.document_type.value if document.document_type else "document",
            additional_context=document.meta_data.get("description"),
        )

        # Store analysis results in meta_data
        document.meta_data = {
            **document.meta_data,
            "analysis": analysis_result.get("data", {}),
            "analysis_status": analysis_result.get("status", "completed"),
            "analyzed_at": datetime.utcnow().isoformat(),
            "content_length_analyzed": analysis_result.get("content_length", 0),
        }

        # Update content if it was extracted
        if content and not document.content:
            document.content = content[:5000]  # Store first 5000 chars

        await db.commit()
        await db.refresh(document)

        logger.info(f"Document analysis completed: {document_id}")

        # Extract analysis data for response
        analysis_data = analysis_result.get("data", {})

        return {
            "success": True,
            "message": "Document analysis completed successfully",
            "document_id": str(document_id),
            "status": "completed",
            "analysis": {
                "project_name": analysis_data.get("project_name"),
                "project_description": analysis_data.get("project_description"),
                "features_count": len(analysis_data.get("features", [])),
                "requirements_count": len(analysis_data.get("requirements", [])),
                "test_scenarios_count": len(analysis_data.get("test_scenarios", [])),
                "complexity": analysis_data.get("complexity"),
                "estimated_effort": analysis_data.get("estimated_effort"),
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to analyze document: {str(e)}"
        )


@router.post("/{document_id}/generate-test-plan")
async def generate_test_plan_from_document(
    document_id: UUID,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Generate a comprehensive test plan from an analyzed document.
    Uses the document analysis results to create a detailed test plan with test suites and cases.
    """
    logger.info(f"=" * 80)
    logger.info(f"📄 GENERATE TEST PLAN FROM DOCUMENT REQUEST")
    logger.info(f"Document ID: {document_id}")
    logger.info(f"User: {current_user.email}")
    logger.info(f"=" * 80)

    result = await db.execute(
        select(DocumentKnowledge).where(DocumentKnowledge.id == document_id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    # Verify project access
    project = await verify_project_access(document.project_id, current_user, db)

    try:
        # Check if document has been analyzed
        analysis_data = document.meta_data.get("analysis")

        if not analysis_data:
            # Analyze document first
            logger.info(f"Document not yet analyzed, analyzing now: {document_id}")

            # Get document content
            file_path = document.meta_data.get("file_path")
            content = ""

            if file_path and os.path.exists(file_path):
                try:
                    import PyPDF2
                    file_ext = Path(file_path).suffix.lower()

                    if file_ext == '.pdf':
                        with open(file_path, 'rb') as f:
                            reader = PyPDF2.PdfReader(f)
                            content_parts = []
                            for page in reader.pages:
                                content_parts.append(page.extract_text())
                            content = '\n'.join(content_parts)
                    elif file_ext in ['.txt', '.md']:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    elif file_ext == '.docx':
                        from docx import Document as DocxDocument
                        doc = DocxDocument(file_path)
                        content = '\n'.join([para.text for para in doc.paragraphs])
                    else:
                        content = document.content or ""
                except Exception as e:
                    logger.warning(f"Error extracting content: {e}")
                    content = document.content or ""
            else:
                content = document.content or ""

            if not content or len(content) < 50:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Document content is too short or empty. Please ensure the document has been uploaded correctly."
                )

            # Analyze the document
            from app.services.document_analysis_service import get_document_analysis_service
            from app.services.ai_service import AIService

            ai_service = AIService()
            analysis_service = get_document_analysis_service(ai_service)

            logger.info(f"🔍 Analyzing document with AI (content length: {len(content)} chars)")
            analysis_result = await analysis_service.analyze_document_content(
                content=content,
                document_type=document.document_type.value if document.document_type else "document",
                additional_context=document.meta_data.get("description"),
            )

            analysis_data = analysis_result.get("data", {})
            logger.info(f"✅ Document analysis completed")
            logger.info(f"   Features found: {len(analysis_data.get('features', []))}")
            logger.info(f"   Requirements found: {len(analysis_data.get('requirements', []))}")
            logger.info(f"   Test scenarios: {len(analysis_data.get('test_scenarios', []))}")

            # Store analysis results
            document.meta_data = {
                **document.meta_data,
                "analysis": analysis_data,
                "analysis_status": analysis_result.get("status", "completed"),
                "analyzed_at": datetime.utcnow().isoformat(),
            }
            await db.commit()

        # Build requirements for comprehensive test plan service
        requirements = {
            "project_id": str(document.project_id),
            "project_type": analysis_data.get("project_type", "web-app"),
            "description": analysis_data.get("project_description", document.document_name),
            "features": [f.get("name") for f in analysis_data.get("features", [])],
            "platforms": analysis_data.get("platforms", ["web"]),
            "priority": "high",  # Default priority
            "complexity": analysis_data.get("complexity", "medium"),
            "timeframe": analysis_data.get("estimated_effort", "2-4 weeks"),
            "objectives": analysis_data.get("objectives", []),
            "requirements": analysis_data.get("requirements", []),
            "test_scenarios": analysis_data.get("test_scenarios", []),
            "constraints": analysis_data.get("constraints", []),
            "assumptions": analysis_data.get("assumptions", []),
            "technical_details": analysis_data.get("technical_details", {}),
        }

        # Generate comprehensive test plan
        from app.services.comprehensive_test_plan_service import get_comprehensive_test_plan_service

        comprehensive_service = get_comprehensive_test_plan_service()
        generation_result = await comprehensive_service.generate_comprehensive_test_plan(
            project_id=document.project_id,
            requirements=requirements,
            db=db,
        )

        if generation_result["status"] != "success":
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate test plan: {generation_result.get('error', 'Unknown error')}",
            )

        # Create test plan in database
        from app.models.test_plan import TestPlan, GenerationType, TestPlanType, ReviewStatus
        from app.models.test_suite import TestSuite
        from app.models.test_case import TestCase, TestCasePriority, TestCaseStatus

        plan_data = generation_result["data"]

        test_plan = TestPlan(
            project_id=document.project_id,
            name=plan_data.get("name", f"Test Plan - {document.document_name}"),
            description=plan_data.get("description"),
            test_plan_type=TestPlanType.REGRESSION.value,

            # IEEE 829 comprehensive sections
            test_objectives_ieee=plan_data.get("test_objectives", []),
            scope_of_testing_ieee=plan_data.get("scope_of_testing", {}),
            test_approach_ieee=plan_data.get("test_approach", {}),
            assumptions_constraints_ieee=plan_data.get("assumptions_and_constraints", []),
            test_schedule_ieee=plan_data.get("test_schedule", {}),
            resources_roles_ieee=plan_data.get("resources_and_roles", []),
            test_environment_ieee=plan_data.get("test_environment", {}),
            entry_exit_criteria_ieee=plan_data.get("entry_exit_criteria", {}),
            risk_management_ieee=plan_data.get("risk_management", {}),
            deliverables_reporting_ieee=plan_data.get("deliverables_and_reporting", {}),
            approval_signoff_ieee=plan_data.get("approval_signoff", {}),

            # Legacy fields
            objectives=[
                obj.get("objective", "") if isinstance(obj, dict) else str(obj)
                for obj in plan_data.get("test_objectives", [])
            ],
            scope_in=plan_data.get("scope_of_testing", {}).get("in_scope", []),
            scope_out=plan_data.get("scope_of_testing", {}).get("out_of_scope", []),

            # Metadata
            generated_by=GenerationType.AI.value,
            source_documents=[str(document_id)],
            confidence_score=str(generation_result.get("confidence", "high")),  # Convert to string
            review_status=ReviewStatus.DRAFT.value,
            tags=plan_data.get("tags", []),
            meta_data={
                "source_document_id": str(document_id),
                "source_document_name": document.document_name,
                "generated_from_analysis": True,
                "estimated_hours": plan_data.get("estimated_hours"),
                "complexity": plan_data.get("complexity"),
                "timeframe": plan_data.get("timeframe"),
            },
            created_by=current_user.email,
        )

        db.add(test_plan)
        await db.flush()

        # Create test suites and test cases
        test_suites_data = plan_data.get("test_suites", [])
        suites_created = 0
        cases_created = 0

        for suite_data in test_suites_data:
            suite_meta_data = suite_data.get("meta_data", {})
            if "category" in suite_data:
                suite_meta_data["category"] = suite_data["category"]

            test_suite = TestSuite(
                project_id=document.project_id,
                test_plan_id=test_plan.id,
                name=suite_data.get("name", "Test Suite"),
                description=suite_data.get("description", ""),
                tags=suite_data.get("tags", []),
                meta_data=suite_meta_data,
                generated_by=GenerationType.AI.value,
                created_by=current_user.email,
            )
            db.add(test_suite)
            await db.flush()
            suites_created += 1

            # Create test cases
            test_cases_data = suite_data.get("test_cases", [])
            for case_data in test_cases_data:
                steps = case_data.get("steps", [])
                if steps and isinstance(steps[0], dict):
                    pass
                else:
                    steps = [
                        {
                            "step_number": idx + 1,
                            "action": step if isinstance(step, str) else step.get("action", ""),
                            "expected_result": step.get("expected_result", "") if isinstance(step, dict) else ""
                        }
                        for idx, step in enumerate(steps)
                    ]

                case_meta_data = case_data.get("meta_data", {})
                if "estimated_time" in case_data:
                    case_meta_data["estimated_time"] = case_data["estimated_time"]

                test_case = TestCase(
                    project_id=document.project_id,
                    test_suite_id=test_suite.id,
                    title=case_data.get("name", "Test Case"),
                    description=case_data.get("description", ""),
                    steps=steps,
                    expected_result=case_data.get("expected_result", ""),
                    priority=TestCasePriority[case_data.get("priority", "medium").upper()].value,
                    status=TestCaseStatus.DRAFT.value,
                    generated_by=GenerationType.AI.value,
                    meta_data=case_meta_data,
                    created_by=current_user.email,
                )
                db.add(test_case)
                cases_created += 1

        await db.commit()
        await db.refresh(test_plan)

        logger.info(f"Test plan generated from document {document_id}: {test_plan.id} with {suites_created} suites and {cases_created} cases")

        return {
            "success": True,
            "message": "Test plan generated successfully from document",
            "test_plan_id": str(test_plan.id),
            "document_id": str(document_id),
            "status": "completed",
            "summary": {
                "test_plan_name": test_plan.name,
                "test_suites_count": suites_created,
                "test_cases_count": cases_created,
                "confidence_score": test_plan.confidence_score,
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating test plan from document: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate test plan: {str(e)}"
        )
